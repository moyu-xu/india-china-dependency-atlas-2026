from __future__ import annotations

import argparse
import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
APP = ROOT / "src" / "App.tsx"
OUT_DIR = ROOT / "public" / "reports"
REPORT_DATE = "2026-07-24"
SNAPSHOT_DATE = "2026-07-24"
COMTRADE = "UN Comtrade 2025 / monthly API, HS 2022 (H6)"
BATTERY_CUSTOMS = json.loads((ROOT / "src" / "data" / "batteryChinaCustoms.json").read_text(encoding="utf-8"))
CHINA_CUSTOMS_HS8 = json.loads((ROOT / "src" / "data" / "chinaCustomsHs8.json").read_text(encoding="utf-8"))
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def block_after(source: str, marker: str) -> str:
    start = source.index(marker)
    start = source.index("=", start)
    start = source.index("[", start)
    depth = 0
    for index in range(start, len(source)):
        char = source[index]
        if char == "[":
            depth += 1
        elif char == "]":
            depth -= 1
            if depth == 0:
                return source[start + 1 : index]
    raise ValueError(f"Cannot find block for {marker}")


def object_after(source: str, marker: str) -> str:
    start = source.index(marker)
    start = source.index("{", start)
    depth = 0
    for index in range(start, len(source)):
        char = source[index]
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return source[start + 1 : index]
    raise ValueError(f"Cannot find object for {marker}")


def split_top_level_entries(block: str) -> list[str]:
    entries: list[str] = []
    depth = 0
    start: int | None = None
    for index, char in enumerate(block):
        if char == "{":
            if depth == 0:
                start = index
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0 and start is not None:
                entries.append(block[start : index + 1])
                start = None
    return entries


def strings_from_array(block: str, field: str) -> list[str]:
    match = re.search(rf"{field}:\s*\[(.*?)\]", block, re.S)
    if not match:
        return []
    return [value.replace('\\"', '"') for value in re.findall(r'"((?:[^"\\]|\\.)*)"', match.group(1))]


def string_field(block: str, field: str, default: str = "") -> str:
    match = re.search(rf"{field}:\s*\"((?:[^\"\\]|\\.)*)\"", block, re.S)
    return match.group(1).replace('\\"', '"') if match else default


def to_hs8(hs: str) -> str:
    return hs


def stat_level(record: dict) -> str:
    hs = record["hs"]
    return "中国 HS8" if len(hs) == 8 else "HS6 国际可比口径"


def money_billions(value: float) -> str:
    if value >= 1:
        return f"{value:.2f} 十亿美元"
    return f"{value * 1000:.1f} 百万美元"


def rmb_billions(value: int) -> str:
    return f"{value / 1_000_000_000:.2f} 十亿元人民币"


def parse_records(source: str) -> list[dict]:
    blocks = [
        block_after(source, "const commodities: CommodityRecord[]"),
        block_after(source, "const fertilizerSubitems: CommodityRecord[]"),
        block_after(source, "const tunnelSubitems: CommodityRecord[]"),
        block_after(source, "const earthmovingSubitems: CommodityRecord[]"),
    ]
    records: list[dict] = []
    pattern = re.compile(
        r'\{\s*id:\s*"(?P<id>[^"]+)".*?hs:\s*"(?P<hs>[^"]+)".*?name:\s*"(?P<name>[^"]+)".*?english:\s*"(?P<english>[^"]+)".*?category:\s*"(?P<category>[^"]+)".*?annual\((?P<china>[\d.]+),\s*(?P<world>[\d.]+)\).*?alternatives:\s*\[(?P<alts>[^\]]*)\].*?definition:\s*"(?P<definition>[^"]*)"',
        re.S,
    )
    for block in blocks:
        for entry in split_top_level_entries(block):
            match = pattern.search(entry)
            if not match:
                continue
            record = match.groupdict()
            record["china"] = float(record["china"])
            record["world"] = float(record["world"])
            record["share"] = record["china"] / record["world"] * 100 if record["world"] else 0
            record["alternatives"] = re.findall(r'"([^"]+)"', record.pop("alts"))
            record["hs8"] = to_hs8(record["hs"])
            record["stat_level"] = stat_level(record)
            record["proxy"] = "proxy: true" in entry
            records.append(record)
    return [record for record in records if record["id"] not in {"fertilizer", "tunnel", "earthmoving"}]


def parse_reports(source: str) -> dict[str, dict]:
    reports_block = object_after(source, "const commodityReports")
    reports: dict[str, dict] = {}
    for match in re.finditer(r"\n\s*([a-zA-Z0-9_]+):\s*\{", reports_block):
        report_id = match.group(1)
        start = reports_block.index("{", match.start())
        depth = 0
        end = start
        for index in range(start, len(reports_block)):
            char = reports_block[index]
            if char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0:
                    end = index + 1
                    break
        block = reports_block[start:end]
        reports[report_id] = {
            "title": string_field(block, "title"),
            "evidence": string_field(block, "evidence", "中等"),
            "status": string_field(block, "status", "公开来源审慎判读"),
            "executive": string_field(block, "executive"),
            "data_points": strings_from_array(block, "dataPoints"),
            "analysis": strings_from_array(block, "analysis"),
            "conclusion": string_field(block, "conclusion"),
            "monitoring": strings_from_array(block, "monitoring"),
            "references": strings_from_array(block, "references"),
            "route_boundary": string_field(block, "routeBoundary", "第三国路径仅用于筛查，不构成转口事实认定。"),
        }
    return reports


def parse_accuracy(source: str) -> dict[str, dict]:
    block = object_after(source, "const reportAccuracyById")
    accuracy: dict[str, dict] = {}
    for report_id, level, reason in re.findall(r'([a-zA-Z0-9_]+):\s*\{\s*level:"([^"]+)",\s*reason:"([^"]+)"\s*\}', block):
        accuracy[report_id] = {"level": level, "reason": reason}
    return accuracy


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc: Document, title: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "中国-印度供应链依赖图谱"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], 9, False, "666666")

    footer = section.footer.paragraphs[0]
    footer.text = f"{title} · {REPORT_DATE}"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], 9, False, "666666")


def add_title(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_run_font(run, 22, True, "000000")
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(16)
    run = sub.add_run(subtitle)
    set_run_font(run, 11, False, "555555")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "000000"):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.text = ""
    run = paragraph.add_run(text)
    set_run_font(run, 9.5, bold, color)


def set_table_geometry(table, widths: list[int]):
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:tblCellMar"):
        for element in list(tbl_pr.findall(qn(tag))):
            tbl_pr.remove(element)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    margins = OxmlElement("w:tblCellMar")
    for edge, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def add_metadata_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for row_index, (label, value) in enumerate(rows):
        cells = table.rows[row_index].cells
        shade_cell(cells[0], "F2F4F7")
        set_cell_text(cells[0], label, True, "1F4D78")
        set_cell_text(cells[1], value)
    set_table_geometry(table, [2300, 7060])
    doc.add_paragraph()


def add_bullets(doc: Document, values: list[str]):
    for value in values:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(value)
        set_run_font(run, 10.5)


def add_source_marker(paragraph, index: int):
    run = paragraph.add_run(f"〔{index}〕")
    set_run_font(run, 8, False, "1F4D78")
    run.font.superscript = True


def add_paragraph_with_source(doc: Document, text: str, footnotes: list[str], source: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, 10.5)
    footnotes.append(source)
    add_source_marker(paragraph, len(footnotes))
    return paragraph


def add_bullets_with_source(doc: Document, values: list[str], footnotes: list[str], source: str):
    for value in values:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(value)
        set_run_font(run, 10.5)
        footnotes.append(source)
        add_source_marker(paragraph, len(footnotes))


def add_bullets_with_sources(doc: Document, values: list[str], footnotes: list[str], sources: list[str]):
    fallback = sources[-1] if sources else "网页内置资料与公开来源整理。"
    for index, value in enumerate(values):
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(value)
        set_run_font(run, 10.5)
        footnotes.append(sources[index] if index < len(sources) else fallback)
        add_source_marker(paragraph, len(footnotes))


def add_footnotes_section(doc: Document, footnotes: list[str]):
    doc.add_heading("脚注与数据来源", level=1)
    if not footnotes:
        doc.add_paragraph("本报告未使用需要单独脚注的数据项。")
        return
    for index, note in enumerate(footnotes, 1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(f"〔{index}〕")
        set_run_font(run, 9, True, "1F4D78")
        run = paragraph.add_run(note)
        set_run_font(run, 9)


def customs_profile(record: dict) -> dict | None:
    return CHINA_CUSTOMS_HS8.get("byHs6", {}).get(record["hs"][:6])


def usd_compact(value: float) -> str:
    if value >= 1_000_000_000:
        return f"{value / 1_000_000_000:.2f} 十亿美元"
    if value >= 1_000_000:
        return f"{value / 1_000_000:.2f} 百万美元"
    return f"{value:,.0f} 美元"


def primary_trade_source(record: dict) -> str:
    return (
        f"UN Comtrade 公共 API，报告国印度（reporterCode=699），"
        f"贸易伙伴中国/世界（partnerCode=156/0），flowCode=M，cmdCode={record['hs']}，period=2025；"
        f"访问日期 {SNAPSHOT_DATE}。"
    )


def customs_source_note(record: dict) -> str:
    profile = customs_profile(record)
    if not profile:
        return (
            f"{CHINA_CUSTOMS_HS8['sourceLabel']}，{CHINA_CUSTOMS_HS8['flow']}，"
            f"未在当前导出文件中匹配到 HS6 {record['hs'][:6]} 的 HS8 明细；"
            f"访问日期 {CHINA_CUSTOMS_HS8['accessedAt']}。"
        )
    return (
        f"{CHINA_CUSTOMS_HS8['sourceLabel']}，{CHINA_CUSTOMS_HS8['flow']}，"
        f"贸易伙伴编码 {CHINA_CUSTOMS_HS8['partnerCode']}，HS6 {record['hs'][:6]} 下真实 HS8 明细，"
        f"来源文件 {'、'.join(CHINA_CUSTOMS_HS8['sourceFiles'])}，访问日期 {CHINA_CUSTOMS_HS8['accessedAt']}。"
    )


def dependence_implication(record: dict, concentration: str) -> str:
    if record["share"] >= 75:
        return (
            f"{record['name']}的对华来源占比已达到高度集中水平。对进口商而言，风险不只体现为单一年度金额占比，"
            "还体现为认证替换、规格兼容、交付周期、价格谈判和售后技术支持的复合约束。"
            "如果企业业务数据继续显示关键型号、核心零部件或上游材料集中在中国供应商，"
            "则全国贸易占比可能低估实际运营依赖。"
        )
    if record["share"] >= 50:
        return (
            f"{record['name']}的对华来源占比较高，说明印度市场在同类商品上已经形成较强的中国供给惯性。"
            "此类商品通常仍存在替代供应国，但替代是否可行取决于价格、技术规格、交货周期、认证资质和客户既有设备体系，"
            "不能只用其他来源国排名判断可替代性。"
        )
    if record["share"] >= 20:
        return (
            f"{record['name']}处于中等对华暴露区间。该区间的政策含义是：总体上并非完全依赖中国，"
            "但中国供应商可能在特定规格、价格带、交付节奏或售后环节形成局部优势。"
            "因此，应优先识别企业采购中是否存在若干关键料号高度集中，而不是简单按全国占比下结论。"
        )
    return (
        f"{record['name']}的公开统计占比处于较低区间。低占比并不意味着没有供应链风险，"
        "但风险更可能集中在少数企业、少数高端型号或受控参数，而不是全国进口来源结构。"
        "后续核验应以企业订单、型号和用途为核心，而不是扩大为行业性依赖判断。"
    )


def policy_implication(record: dict) -> str:
    return (
        f"从中国出口管制策略研究角度看，HS {record['hs']} 只能作为筛查入口，不能直接替代受控物项判定。"
        "真正需要落到三个层级：第一，税号与商品描述是否锁定到具体物项；第二，技术参数、性能指标、材料形态或最终用途是否触及管制清单；"
        "第三，交易链条中的最终用户、最终用途、发票国、装运国和原产国是否一致。"
        "因此，本报告将贸易金额用于识别优先级，将单证和技术参数留作后续核验重点。"
    )


def substitution_note(record: dict) -> str:
    alternatives = "、".join(record["alternatives"]) if record["alternatives"] else "公开统计未显示稳定替代来源"
    return (
        f"替代供应方面，2025 年除中国外的主要来源包括：{alternatives}。"
        "这些来源国可以作为供应分散化的第一批观察对象，但其是否具备现实替代能力，仍需进一步比较产品等级、报价条款、交期、认证记录、"
        "历史交付稳定性和售后能力。对于工程设备、电子电力和医药化工类商品，还需要特别注意整机、零部件、材料和中间体之间的税号差异。"
    )


def commodity_report_doc(record: dict, report: dict, accuracy: dict) -> Document:
    title = f"{record['name']}对华进口依赖与供应风险分析报告"
    footnotes: list[str] = []
    doc = Document()
    style_document(doc, title)
    add_title(doc, title, f"商品：{record['name']} / {record['english']} · 快照 {SNAPSHOT_DATE}")
    add_metadata_table(
        doc,
        [
            ("真实统计编码", f"{record['stat_level']} {record['hs8']}"),
            ("统计口径", record["stat_level"]),
            ("2025 自中国进口", money_billions(record["china"])),
            ("2025 全球进口", money_billions(record["world"])),
            ("对华来源占比", f"{record['share']:.1f}%"),
            ("结论准确度", f"{accuracy.get('level', '高概率')}：{accuracy.get('reason', '结论基于公开统计与审慎边界。')}"),
        ],
    )

    doc.add_heading("一、核心判断", level=1)
    concentration = "高度集中" if record["share"] >= 75 else "较高" if record["share"] >= 50 else "中等" if record["share"] >= 20 else "较低"
    add_paragraph_with_source(
        doc,
        f"2025 年，印度进口“{record['name']}”中来自中国的金额占比为 {record['share']:.1f}%，"
        f"按本报告阈值属于{concentration}。该判断严格限定于 HS 2022 六位商品编码 {record['hs']}，"
        "不外推至父级税目、具体企业或受控技术参数。",
        footnotes,
        primary_trade_source(record),
    )
    if report.get("executive"):
        add_paragraph_with_source(doc, report["executive"], footnotes, primary_trade_source(record))
    doc.add_paragraph(
        "本报告的用途是为业务数据复核、供应链审计和出口管制策略研究提供公开来源底稿。"
        "因此，报告将“贸易统计能够证明的来源结构”和“仍需单证核验的路径推演”分开表述：前者用于判断依赖程度，后者只用于形成核验清单。"
    )

    doc.add_heading("二、商品边界与统计口径", level=1)
    doc.add_paragraph(record["definition"] or f"本报告以 HS {record['hs']} 作为统计边界，商品名称为“{record['name']}”。")
    doc.add_paragraph(
        f"网页矩阵和本报告均优先使用真实可复核编码。当前国际比较口径为 {record['stat_level']} {record['hs8']}；"
        "如公开来源只能稳定提供 HS6，本报告即如实使用 HS6，不补零、不以父级税目金额替代具体物项。"
    )

    doc.add_heading("三、数据事实", level=1)
    data_points = report.get("data_points") or [
        f"2025 年印度自中国进口 {money_billions(record['china'])}，全球进口 {money_billions(record['world'])}，对华来源占比 {record['share']:.1f}%。",
        f"页面和报告均使用真实的 {record['stat_level']} {record['hs8']}；不补零、不使用父级大类代理金额。",
        f"剔除中国后，2025 年其他主要供应来源为：{'、'.join(record['alternatives']) if record['alternatives'] else '未报告其他境外来源'}。",
    ]
    report_refs = report.get("references") or []
    data_sources = [primary_trade_source(record)]
    if report_refs:
        data_sources.extend(report_refs[1:] if len(report_refs) > 1 else report_refs)
    add_bullets_with_sources(doc, data_points, footnotes, data_sources)
    doc.add_paragraph(
        "上述金额均为现价美元口径。印度进口统计通常采用 CIF 进口口径；中国海关出口镜像采用中国出口统计口径。"
        "两类口径可以互相校验方向和异常，但不能直接相加或替代。"
    )

    profile = customs_profile(record)
    doc.add_heading("四、中国海关 HS8 出口镜像", level=1)
    if profile:
        annual_2025 = profile.get("annual", {}).get("2025", {}).get("usd", 0)
        annual_2026 = profile.get("annual", {}).get("2026", {}).get("usd", 0)
        hs8_count = len(profile.get("hs8", []))
        months = [m for m in profile.get("months", []) if "202501" <= m.get("period", "") <= "202606"]
        peak = max(months, key=lambda item: item.get("usd", 0), default={"period": "—", "usd": 0})
        add_paragraph_with_source(
            doc,
            f"在中国海关出口镜像中，HS6 {record['hs'][:6]} 下共匹配 {hs8_count} 个真实 HS8 子项。"
            f"2025 年中国出口至印度金额为 {usd_compact(annual_2025)}；2026 年已导出月份合计为 {usd_compact(annual_2026)}。"
            f"2025—2026 已导出月份中，月度峰值出现在 {peak.get('period')}，金额为 {usd_compact(peak.get('usd', 0))}。",
            footnotes,
            customs_source_note(record),
        )
        if profile.get("hs8"):
            table = doc.add_table(rows=1, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            for index, header in enumerate(["HS8", "商品名称", "2025出口额", "主要贸易方式"]):
                shade_cell(table.rows[0].cells[index], "F2F4F7")
                set_cell_text(table.rows[0].cells[index], header, True, "1F4D78")
            repeat_table_header(table.rows[0])
            for item in profile.get("hs8", [])[:10]:
                row = table.add_row().cells
                values = [
                    item.get("code", ""),
                    item.get("name", ""),
                    usd_compact(item.get("annual", {}).get("2025", {}).get("usd", 0)),
                    (item.get("tradeModes") or [{}])[0].get("name", "—"),
                ]
                for index, value in enumerate(values):
                    set_cell_text(row[index], value)
            set_table_geometry(table, [1300, 4260, 1700, 2100])
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("表注：如某 HS6 下超过 10 个 HS8，表中列示金额靠前的主要子项；完整明细以网页交互卡片和原始 CSV 为准。")
            set_run_font(run, 9, False, "666666")
    else:
        add_paragraph_with_source(
            doc,
            f"当前中国海关 HS8 导出文件尚未覆盖 HS6 {record['hs'][:6]}。因此，本商品报告仍以印度报告进口口径判断依赖率，"
            "中国出口镜像部分留待后续补充 CSV 后更新。",
            footnotes,
            customs_source_note(record),
        )

    doc.add_heading("五、第三国路径与判读边界", level=1)
    doc.add_paragraph(report.get("route_boundary") or "第三国路径仅用于筛查，不构成转口事实认定。")
    doc.add_paragraph(
        f"对 HS {record['hs']} 而言，第三国路径至少需要同时满足四类证据：第一，中国至中转节点存在同口径商品出口；"
        "第二，中转节点至印度存在同口径商品出口或印度进口；第三，金额、月份、数量单位和商品描述具有可比性；"
        "第四，原产地证书、提单、发票、加工记录或公开查发案例能够解释同一货物或同一供应网络。"
    )

    doc.add_heading("六、分析", level=1)
    analysis = report.get("analysis") or [
        f"印度自中国进口 {money_billions(record['china'])}，相对于该商品进口总额 {money_billions(record['world'])}；比例使用未舍入金额计算，页面显示值仅作四舍五入。",
        "其他供应来源按同一 HS6、同一报告国、同一年度的进口金额排序。它们说明采购来源结构，不等于已经具备同等产能、认证、交付周期或价格条件。",
        "公开贸易数据能够识别来源集中度，但不足以替代企业级 BOM、合同、发票、原产地证书和物流单证。对受控物项还必须核对技术参数、最终用户和最终用途。",
    ]
    for paragraph in analysis:
        doc.add_paragraph(paragraph)
    doc.add_paragraph(dependence_implication(record, concentration))
    doc.add_paragraph(substitution_note(record))
    doc.add_paragraph(policy_implication(record))

    doc.add_heading("七、结论", level=1)
    conclusion = report.get("conclusion") or (
        f"{record['name']}对华进口来源集中度{concentration}，应"
        f"{'列入优先核验清单' if record['share'] >= 50 else '保持常态监测'}。"
        f"本结论准确度标注为“{accuracy.get('level', '高概率')}”，仅适用于 HS6 {record['hs']} 的 2025 年进口来源结构；"
        "后续若取得中国海关 HS8、企业料号和业务单证，应在不改变国际 HS6 可比口径的前提下进一步下钻。"
    )
    doc.add_paragraph(conclusion)
    doc.add_paragraph(
        f"准确度标注：{accuracy.get('level', '高概率')}。"
        f"{accuracy.get('reason', '结论主要基于公开统计与审慎边界表述。')}"
    )

    doc.add_heading("八、后续监测重点", level=1)
    monitoring = report.get("monitoring") or [
        f"HS6 {record['hs']} 的月度进口金额、数量和对华来源占比",
        "中国海关 HS8、企业料号与印度本国八位税号映射",
        "原产国、装运国、发票国、提单路径与实际加工工序",
        "出口许可证、技术参数、最终用户和最终用途变化",
    ]
    add_bullets(doc, monitoring)

    references = report.get("references") or []
    for reference in references:
        footnotes.append(reference)
    footnotes.extend([
        primary_trade_source(record),
        "印度 DGCI&S TradeStat：月度贸易数据库与编码调整说明。",
        "中国商务部、海关总署及中国出口管制信息网：现行两用物项清单、公告与许可证管理目录。",
        customs_source_note(record),
    ])
    add_footnotes_section(doc, footnotes)
    doc.add_paragraph("注：本报告用于公开来源研究和合规初筛，不构成法律意见或个案事实认定。")
    return doc


def overall_report_doc(records: list[dict]) -> Document:
    doc = Document()
    footnotes: list[str] = []
    style_document(doc, "总分析报告")
    add_title(doc, "中国-印度供应链依赖图谱总分析报告", f"重点商品矩阵、化肥专题、工程设备专题与报告下载索引 · 快照 {SNAPSHOT_DATE}")
    china_total = sum(record["china"] for record in records)
    world_total = sum(record["world"] for record in records)
    high_records = [record for record in records if record["share"] >= 50]
    mid_records = [record for record in records if 20 <= record["share"] < 50]
    add_metadata_table(
        doc,
        [
            ("报告范围", "重点商品矩阵及可交互子项"),
            ("分类版本", "HS 2022；按公开来源最细可核验层级显示真实 HS6/HS8"),
            ("统计边界", "仅汇总互不重叠的具体商品物项；不使用补零代理或父级大类重复金额"),
            ("加权对华来源占比", f"{china_total / world_total * 100:.1f}%"),
            ("研究日期", REPORT_DATE),
        ],
    )
    doc.add_heading("一、总览结论", level=1)
    add_paragraph_with_source(
        doc,
        f"本报告汇总 {len(records)} 个互不重叠的 HS 2022 六位商品物项。2025 年这些商品自中国进口合计"
        f" {money_billions(china_total)}，商品进口总额合计 {money_billions(world_total)}，加权对华来源占比为"
        f" {china_total / world_total * 100:.1f}%。矩阵不计入 HS31、HS4 或专题父级金额，也不使用补零八位码。",
        footnotes,
        f"UN Comtrade 公共 API：印度报告进口，HS 2022 六位商品口径，period=2025；访问 {SNAPSHOT_DATE}。",
    )
    doc.add_paragraph(
        f"从结构看，{len(high_records)} 个商品的对华来源占比达到 50% 以上，属于优先核验区；"
        f"{len(mid_records)} 个商品处于 20%—50% 的中等暴露区间，适合作为月度监测和业务抽样复核对象。"
        "这些商品并不代表印度全口径进口依赖，而是围绕原材料、医药化工、电子电力、工程设备、车辆零部件等重点物项构建的风险样本。"
    )
    doc.add_paragraph(
        "本报告的核心结论是：对华依赖不能只看父级大类金额，也不能把 HS4、HS6 或补零八位码混为一谈。"
        "政策研究应尽量落到真实商品物项，再结合中国海关 HS8、印度进口口径、企业料号和单证链条进行交叉验证。"
        "在此基础上，第三国路径才有可能从“统计信号”升级为“可核验路径”。"
    )
    doc.add_paragraph(
        "第三国路径应作为筛查信号而非事实认定。报告中的多节点路径允许出现两至三个中转经济体，"
        "但只有在金额、时间、产品、原产地和物流单证闭合后，才能进入个案判断。"
    )

    doc.add_heading("二、重点商品表", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["具体商品", "统计编码", "统计口径", "从中国进口", "商品进口总额", "从中国进口占比"]
    for index, header in enumerate(headers):
        shade_cell(table.rows[0].cells[index], "F2F4F7")
        set_cell_text(table.rows[0].cells[index], header, True, "1F4D78")
    repeat_table_header(table.rows[0])
    for record in sorted(records, key=lambda item: item["share"], reverse=True):
        row = table.add_row().cells
        values = [
            record["name"],
            record["hs8"],
            record["stat_level"],
            money_billions(record["china"]),
            money_billions(record["world"]),
            f"{record['share']:.1f}%",
        ]
        for index, value in enumerate(values):
            set_cell_text(row[index], value)
    set_table_geometry(table, [2200, 1080, 1480, 1500, 1500, 1600])

    doc.add_heading("三、重点板块分析", level=1)
    doc.add_heading("3.1 电子与电力设备", level=2)
    doc.add_paragraph(
        "电子与电力设备是样本中金额和集中度最突出的板块。锂离子蓄电池、半导体器件、集成电路和变压器/静止式变流器分别对应电芯与储能、光伏与半导体、电控核心部件以及电力电子基础设施。"
        "这类商品的共同特点是：贸易金额能说明来源集中度，但不能直接说明受控参数；真正的政策判断需要进一步识别晶圆/芯片型号、电芯化学体系、功率等级、BMS、逆变器拓扑和最终应用场景。"
    )
    doc.add_heading("3.2 化肥专题", level=2)
    doc.add_paragraph(
        "化肥应按总项和子项分开。HS31 总项用于观察整体暴露；尿素、DAP、MOP、NPK 分别对应不同采购周期、替代来源和政策敏感性，财年数量与自然年价值不可直接混算。"
        "在政策层面，化肥更接近民生和粮食安全品类，其供应风险评估应优先关注年度采购安排、政府补贴、长期协议和港口到货节奏，而不是简单套用工业品依赖模型。"
    )
    doc.add_heading("3.3 工程设备与车辆", level=2)
    doc.add_paragraph(
        "工程设备专题以 HS 843031、843039、870410、870510、870540 和 843143 等法定六位商品物项分别统计。专题标题仅作导航；其中 HS 843031/843039 还包含采煤机和截岩机，不能把税目金额直接称为盾构机成交额或台数。"
        "工程车和工程机械的公开统计容易受到项目制交付、临时进口、整机/部件拆分、FOB/CIF 镜像差异和数量单位差异影响，因此镜像贸易差异是审计触发器，不是转口证明。"
    )
    doc.add_heading("3.4 原材料、医药化工与机械零部件", level=2)
    doc.add_paragraph(
        "天然石墨、稀土化合物、医药中间体、泵阀、机床附件和汽车零部件等商品的政策含义更偏向“参数核验”和“供应链黏性”。"
        "其中，石墨和稀土的 HS6 金额无法替代纯度、形态、粒径、元素组成和最终用途判断；医药化工品的替代难点往往在注册、质量体系和长期供货稳定性；机械零部件则需要结合设备存量、接口标准和维保体系。"
    )
    doc.add_heading("四、第三国路径与交叉验证方法", level=1)
    add_paragraph_with_source(
        doc,
        "第三国路径模块采用分段交叉验证：第一段优先使用中国海关出口统计，观察中国至中转国的真实出口；后续段使用公开报告国数据或印度进口端数据观察中转国至印度的贸易信号。"
        "如果存在多个中转国，必须逐段验证，不能只因最终印度进口增加就反推完整路径。",
        footnotes,
        f"{CHINA_CUSTOMS_HS8['sourceLabel']}与 UN Comtrade 公共 API；中国海关数据访问 {CHINA_CUSTOMS_HS8['accessedAt']}，UN Comtrade 访问 {SNAPSHOT_DATE}。",
    )
    doc.add_paragraph(
        "本报告将路径证据分为三层：第一层是同一 HS6/HS8 在相邻贸易段同时出现；第二层是金额、月份、数量单位和商品描述具有可比性；第三层是提单、发票、原产地证书、加工记录或公开查发案例能够确认同一货物或同一供应网络。"
        "只有第三层证据成立，才可以进入个案层面的转口事实判断。"
    )
    doc.add_heading("五、对中国出口管制策略研究的启示", level=1)
    doc.add_paragraph(
        "若研究目标是中国出口管制策略，HS 编码应作为入口而非结论。建议先用 HS6 建立国际可比样本，再用中国 HS8、监管证件、两用物项清单、企业产品参数和最终用户信息下钻。"
        "对于高占比商品，应优先核验是否存在可控参数或关键技术环节；对于低占比但政策敏感商品，应避免被全国金额低估误导，转而关注企业级采购和具体型号。"
    )
    doc.add_paragraph(
        "在执法和合规层面，最有价值的并不是简单列出“中转国名单”，而是形成可复核的核验清单：哪些 HS8 子项金额异常、哪些贸易方式集中、哪些月份出现峰值、哪些第三国同时承担进口和对印出口、哪些公开案例显示类似规避路径。"
        "这类清单可以和企业订单、船运单证、原产地证书、付款路径和最终用途声明结合，构成后续业务验证框架。"
    )
    doc.add_heading("六、方法与限制", level=1)
    add_bullets(
        doc,
        [
            "依赖率 = 印度自中国进口额 ÷ 印度全球进口额。",
            "公开统计统一使用真实 HS 2022 六位商品编码；不补零、不使用父级大类或代理金额。中国 HS8 仅在取得海关或企业业务数据后用于向下映射。",
            f"自动审计逐项核对 {len(records)} 个商品的 2025 年世界/中国进口额、12 个月月度合计及其他供应国排序，并使用缓存、限速和 429 自动重试。",
            "结论准确度分为高概率、低概率、推测，表示公开证据强弱，不代表法律结论。",
            "本报告不对违法转口、规避关税或规避出口管制作事实认定。",
        ],
    )
    add_footnotes_section(doc, footnotes)
    return doc


def main():
    parser = argparse.ArgumentParser(description="Generate website Word reports")
    parser.add_argument("--only", action="append", choices=["overall", "battery"], help="Regenerate only the selected report")
    args = parser.parse_args()
    source = APP.read_text(encoding="utf-8")
    records = parse_records(source)
    reports = parse_reports(source)
    accuracy = parse_accuracy(source)
    selected = set(args.only or [])
    if not selected and OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    generated = 0
    if not selected or "overall" in selected:
        overall = overall_report_doc(records)
        overall.save(OUT_DIR / "overall.docx")
        generated += 1

    for record in records:
        if selected and record["id"] not in selected:
            continue
        doc = commodity_report_doc(record, reports.get(record["id"], {}), accuracy.get(record["id"], {}))
        doc.save(OUT_DIR / f"{record['id']}.docx")
        generated += 1

    print(f"Generated {generated} Word report(s) in {OUT_DIR}")


if __name__ == "__main__":
    main()
