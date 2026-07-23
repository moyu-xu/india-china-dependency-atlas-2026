from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
APP = ROOT / "src" / "App.tsx"
REPORT_DIR = ROOT / "public" / "reports"
EXPECTED_COUNT = 22
FORBIDDEN_TEXT = (
    "8 位编码用于后续业务单证对齐",
    "公开统计层级不足时，仍标注 HS31、HS4",
    "含氮/含氧化工品",
    "页面展示 8 位筛查码",
)


def table_geometry(table) -> tuple[int, int, list[int]]:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    grid = table._tbl.tblGrid
    return (
        int(tbl_w.get(qn("w:w"))) if tbl_w is not None else -1,
        int(tbl_ind.get(qn("w:w"))) if tbl_ind is not None else -1,
        [int(column.get(qn("w:w"))) for column in grid.findall(qn("w:gridCol"))],
    )


source = APP.read_text(encoding="utf-8")
ids = set(re.findall(r'\{\s*id:\s*"([a-zA-Z0-9_]+)",\s*hs:', source))
ids -= {"fertilizer", "tunnel", "earthmoving"}
expected_files = {f"{item_id}.docx" for item_id in ids} | {"overall.docx"}
actual_files = {path.name for path in REPORT_DIR.glob("*.docx")}

errors: list[str] = []
if len(actual_files) != EXPECTED_COUNT:
    errors.append(f"Expected {EXPECTED_COUNT} reports, found {len(actual_files)}")
if actual_files != expected_files:
    errors.append(f"Report file mismatch: missing={sorted(expected_files-actual_files)} extra={sorted(actual_files-expected_files)}")

for path in sorted(REPORT_DIR.glob("*.docx")):
    with zipfile.ZipFile(path) as archive:
        damaged = archive.testzip()
        if damaged:
            errors.append(f"{path.name}: corrupt ZIP member {damaged}")
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for forbidden in FORBIDDEN_TEXT:
        if forbidden in text:
            errors.append(f"{path.name}: stale wording found: {forbidden}")
    required = ("数据事实", "结论", "来源") if path.name != "overall.docx" else ("总览结论", "重点商品表", "方法与限制")
    for phrase in required:
        if phrase not in text:
            errors.append(f"{path.name}: missing required section {phrase}")
    if path.name != "overall.docx" and "结论准确度" not in "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    ):
        errors.append(f"{path.name}: missing accuracy label")
    for table_index, table in enumerate(document.tables):
        width, indent, grid = table_geometry(table)
        if width != 9360 or indent != 120 or sum(grid) != 9360:
            errors.append(
                f"{path.name}: table {table_index} geometry width={width} indent={indent} grid={grid}"
            )

summary = {
    "reports": len(actual_files),
    "commodityReports": len(actual_files - {"overall.docx"}),
    "errors": len(errors),
}
print(summary)
if errors:
    print("\n".join(errors))
    raise SystemExit(2)
